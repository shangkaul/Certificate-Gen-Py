from flask import Flask, request,render_template,flash
from werkzeug.utils import secure_filename
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import PyPDF2
import sys
import csv
import os.path

import sys
import subprocess
import re
import time


import csv
import email,smtplib, ssl

from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart



app = Flask(__name__)

# cwd = os.getcwd()
# dir = cwd+'/uploads'
# app.config['UPLOAD_FOLDER']=dir
# print(dir)
app.config["DEBUG"] = True



def send_mail():
    interns=[]
    count=0
    sender_email = "iwinterns19@gmail.com"
    password = "Internships@2020iw"
    with open('mail.csv') as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=',')
        for row in csv_reader:
            interns.append(row)

    for row in interns:
        receiver_email = row[0]
        message = MIMEMultipart()
        message["Subject"] = "Internship Certificate | Certificate Of Internship"
        message["From"] = sender_email
        message["To"] = receiver_email
        message["Bcc"] = sender_email

        
        html = """<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Transitional //EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-transitional.dtd">

<html xmlns="http://www.w3.org/1999/xhtml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:v="urn:schemas-microsoft-com:vml">
<head>
<!--[if gte mso 9]><xml><o:OfficeDocumentSettings><o:AllowPNG/><o:PixelsPerInch>96</o:PixelsPerInch></o:OfficeDocumentSettings></xml><![endif]-->
<meta content="text/html; charset=utf-8" http-equiv="Content-Type"/>
<meta content="width=device-width" name="viewport"/>
<!--[if !mso]><!-->
<meta content="IE=edge" http-equiv="X-UA-Compatible"/>
<!--<![endif]-->
<title></title>
<!--[if !mso]><!-->
<link href="https://fonts.googleapis.com/css?family=Cabin" rel="stylesheet" type="text/css"/>
<link href="https://fonts.googleapis.com/css?family=Merriweather" rel="stylesheet" type="text/css"/>
<link href="https://fonts.googleapis.com/css?family=Oswald" rel="stylesheet" type="text/css"/>
<link href="https://fonts.googleapis.com/css?family=Roboto+Slab" rel="stylesheet" type="text/css"/>
<link href="https://fonts.googleapis.com/css?family=Open+Sans" rel="stylesheet" type="text/css"/>
<link href="https://fonts.googleapis.com/css?family=Lato" rel="stylesheet" type="text/css"/>
<!--<![endif]-->
<style type="text/css">
		body {
			margin: 0;
			padding: 0;
		}
		a:link {
  text-decoration: none;
}

a:visited {
  text-decoration: none;
}

a:hover {
  text-decoration: underline;
}

a:active {
  text-decoration: underline;
}
		table,
		td,
		tr {
			vertical-align: top;
			border-collapse: collapse;
		}

		* {
			line-height: inherit;
		}

		a[x-apple-data-detectors=true] {
			color: inherit !important;
			text-decoration: none !important;
		}
	</style>
<style id="media-query" type="text/css">
		@media (max-width: 660px) {

			.block-grid,
			.col {
				min-width: 320px !important;
				max-width: 100% !important;
				display: block !important;
			}

			.block-grid {
				width: 100% !important;
			}

			.col {
				width: 100% !important;
			}

			.col>div {
				margin: 0 auto;
			}

			img.fullwidth,
			img.fullwidthOnMobile {
				max-width: 100% !important;
			}

			.no-stack .col {
				min-width: 0 !important;
				display: table-cell !important;
			}

			.no-stack.two-up .col {
				width: 50% !important;
			}

			.no-stack .col.num4 {
				width: 33% !important;
			}

			.no-stack .col.num8 {
				width: 66% !important;
			}

			.no-stack .col.num4 {
				width: 33% !important;
			}

			.no-stack .col.num3 {
				width: 25% !important;
			}

			.no-stack .col.num6 {
				width: 50% !important;
			}

			.no-stack .col.num9 {
				width: 75% !important;
			}

			.video-block {
				max-width: none !important;
			}

			.mobile_hide {
				min-height: 0px;
				max-height: 0px;
				max-width: 0px;
				display: none;
				overflow: hidden;
				font-size: 0px;
			}

			.desktop_hide {
				display: block !important;
				max-height: none !important;
			}
			a:link {
  text-decoration: none;
}

a:visited {
  text-decoration: none;
}

a:hover {
  text-decoration: underline;
}

a:active {
  text-decoration: underline;
}
		}

		
	</style>
<style id="menu-media-query" type="text/css">
		@media (max-width: 660px) {
			.menu-checkbox[type="checkbox"]~.menu-links {
				display: none !important;
				padding: 5px 0;
			}

			.menu-checkbox[type="checkbox"]~.menu-links span.sep {
				display: none;
			}

			.menu-checkbox[type="checkbox"]:checked~.menu-links,
			.menu-checkbox[type="checkbox"]~.menu-trigger {
				display: block !important;
				max-width: none !important;
				max-height: none !important;
				font-size: inherit !important;
			}

			.menu-checkbox[type="checkbox"]~.menu-links>a,
			.menu-checkbox[type="checkbox"]~.menu-links>span.label {
				display: block !important;
				text-align: center;
			}

			.menu-checkbox[type="checkbox"]:checked~.menu-trigger .menu-close {
				display: block !important;
			}

			.menu-checkbox[type="checkbox"]:checked~.menu-trigger .menu-open {
				display: none !important;
			}

			#menudnm8m6~div label {
				border-radius: 0% !important;
			}

			#menudnm8m6:checked~.menu-links {
				background-color: #2e9fdd !important;
			}

			#menudnm8m6:checked~.menu-links a {
				color: #ffffff !important;
			}

			#menudnm8m6:checked~.menu-links span {
				color: #ffffff !important;
			}
			a:link {
  text-decoration: none;
}

a:visited {
  text-decoration: none;
}

a:hover {
  text-decoration: underline;
}

a:active {
  text-decoration: underline;
}
		}
	</style>
	<style>
		a:link {
  text-decoration: none;
}

a:visited {
  text-decoration: none;
}

a:hover {
  text-decoration: underline;
}

a:active {
  text-decoration: underline;
}
	</style>
</head>
<body class="clean-body" style="margin: 0; padding: 0; -webkit-text-size-adjust: 100%; background-color: #ffffff;">
	<div style="display: none; max-height: 0px; overflow: hidden;">
		Hi """+row[1]+""",
		  <br>      
				It was a great experience working with you and we wish you all a great luck for your life ahead and may you achieve all your goals and ambitions. We've enclosed your certificate in this email. 
				Team Inception Wave feels grateful to have you as our interns. We seek your support for the future as well. 
				You can follow us on Instagram, Facebook and LinkedIn to show your support!
				Regards,
				Inception Wave Pvt. Ltd.
		</div>
		
		<!-- Insert &zwnj;&nbsp; after hidden preview text -->
		<div style="display: none; max-height: 0px; overflow: hidden;">
		&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;&zwnj;&nbsp;
		</div>
<!--[if IE]><div class="ie-browser"><![endif]-->
<table bgcolor="#ffffff" cellpadding="0" cellspacing="0" class="nl-container" role="presentation" style="table-layout: fixed; vertical-align: top; min-width: 320px; Margin: 0 auto; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; background-color: #ffffff; width: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td style="word-break: break-word; vertical-align: top;" valign="top">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td align="center" style="background-color:#ffffff"><![endif]-->
<div style="background-color:#2e9fdd;">
<div class="block-grid two-up no-stack" style="Margin: 0 auto; min-width: 320px; max-width: 640px; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; background-color: transparent;">
<div style="border-collapse: collapse;display: table;width: 100%;background-color:transparent;">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#2e9fdd;"><tr><td align="center"><table cellpadding="0" cellspacing="0" border="0" style="width:640px"><tr class="layout-full-width" style="background-color:transparent"><![endif]-->
<!--[if (mso)|(IE)]><td align="center" width="320" style="background-color:transparent;width:320px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:5px; padding-bottom:5px;"><![endif]-->
<div class="col num6" style="min-width: 320px; max-width: 320px; display: table-cell; vertical-align: top; width: 320px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:5px; padding-bottom:5px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<div align="left" class="img-container left autowidth" style="padding-right: 0px;padding-left: 20px;">
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr style="line-height:0px"><td style="padding-right: 0px;padding-left: 20px;" align="left"><![endif]-->
<div style="font-size:1px;line-height:20px"> </div><img alt="Inception Wave" border="0" class="left autowidth" src="https://srmpediaesyes.files.wordpress.com/2020/08/logo_iw.png" style="text-decoration: none; -ms-interpolation-mode: bicubic; height: auto; border: 0; width: 100%; max-width: 95px; display: block;" title="Image" width="95"/>
<!--[if mso]></td></tr></table><![endif]-->
</div>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td><td align="center" width="320" style="background-color:transparent;width:320px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:5px; padding-bottom:5px;"><![endif]-->
<div class="col num6" style="min-width: 320px; max-width: 320px; display: table-cell; vertical-align: top; width: 320px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:5px; padding-bottom:5px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<div class="mobile_hide">
<table border="0" cellpadding="0" cellspacing="0" class="divider" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td class="divider_inner" style="word-break: break-word; vertical-align: top; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%; padding-top: 10px; padding-right: 10px; padding-bottom: 10px; padding-left: 10px;" valign="top">
<table align="center" border="0" cellpadding="0" cellspacing="0" class="divider_content" height="0" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; border-top: 0px solid transparent; height: 0px; width: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td height="0" style="word-break: break-word; vertical-align: top; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top"><span></span></td>
</tr>
</tbody>
</table>
</td>
</tr>
</tbody>
</table>
</div>
<table border="0" cellpadding="0" cellspacing="0" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt;" valign="top" width="100%">
<tr style="vertical-align: top;" valign="top">
<td align="center" style="word-break: break-word; vertical-align: top; padding-top: 15px; padding-bottom: 0px; padding-left: 0px; padding-right: 0px; text-align: center; font-size: 0px;" valign="top">
<!--[if !mso><!--> <input class="menu-checkbox" id="menudnm8m6" style="display:none !important;max-height:0;visibility:hidden;" type="checkbox"/>
<!--<![endif]-->
<div class="menu-trigger" style="display:none;max-height:0px;max-width:0px;font-size:0px;overflow:hidden;"> <label class="menu-label" for="menudnm8m6" style="height:48px;width:48px;display:inline-block;cursor:pointer;mso-hide:all;user-select:none;align:center;text-align:center;color:#ffffff;text-decoration:none;background-color:#2e9fdd;"><span class="menu-open" style="mso-hide:all;font-size:38px;line-height:48px;">☰</span><span class="menu-close" style="display:none;mso-hide:all;font-size:38px;line-height:48px;">✕</span></label></div>
<div class="menu-links">
<!--[if mso]>
<table role="presentation" border="0" cellpadding="0" cellspacing="0" align="center">
<tr>
<td style="padding-top:5px;padding-right:5px;padding-bottom:5px;padding-left:5px">
<![endif]-->
<!--[if mso]></td><td><![endif]-->
<!--[if mso]></td><![endif]-->
<!--[if mso]></td><td style="padding-top:5px;padding-right:5px;padding-bottom:5px;padding-left:5px"><![endif]--><span class="label" style="padding-top:5px;padding-bottom:5px;padding-left:5px;padding-right:5px;display:inline;font-family:Merriwheater, Georgia, serif;font-size:13px;color:#000000;"><a style="color: white;" href="https://inceptionwave.in/home">Inception Wave</a></span>
<!--[if mso]></td><td><![endif]--><span class="sep" style="font-size:13px;font-family:Merriwheater, Georgia, serif;color:#000000;">|</span>
<!--[if mso]></td><![endif]-->
<!--[if mso]></td><td style="padding-top:5px;padding-right:5px;padding-bottom:5px;padding-left:5px"><![endif]--><span class="label" style="padding-top:5px;padding-bottom:5px;padding-left:5px;padding-right:5px;display:inline;font-family:Merriwheater, Georgia, serif;font-size:13px;color:#000000;"><a style="color: white;" href="http://grapido.in/home">Grapido</a></span>
<!--[if mso]></td></tr></table><![endif]-->
</div>
</td>
</tr>
</table>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td></tr></table></td></tr></table><![endif]-->
</div>
</div>
</div>
<div style="background-color:#2e9fdd;">
<div class="block-grid" style="Margin: 0 auto; min-width: 320px; max-width: 640px; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; background-color: transparent;">
<div style="border-collapse: collapse;display: table;width: 100%;background-color:transparent;">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#2e9fdd;"><tr><td align="center"><table cellpadding="0" cellspacing="0" border="0" style="width:640px"><tr class="layout-full-width" style="background-color:transparent"><![endif]-->
<!--[if (mso)|(IE)]><td align="center" width="640" style="background-color:transparent;width:640px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:0px; padding-bottom:0px;"><![endif]-->
<div class="col num12" style="min-width: 320px; max-width: 640px; display: table-cell; vertical-align: top; width: 640px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:0px; padding-bottom:0px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<table border="0" cellpadding="0" cellspacing="0" class="divider" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td class="divider_inner" style="word-break: break-word; vertical-align: top; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%; padding-top: 10px; padding-right: 10px; padding-bottom: 10px; padding-left: 10px;" valign="top">
<table align="center" border="0" cellpadding="0" cellspacing="0" class="divider_content" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; border-top: 1px solid #000000; width: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td style="word-break: break-word; vertical-align: top; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top"><span></span></td>
</tr>
</tbody>
</table>
</td>
</tr>
</tbody>
</table>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td></tr></table></td></tr></table><![endif]-->
</div>
</div>
</div>
<div style="background-color:#2e9fdd;">
<div class="block-grid two-up" style="Margin: 0 auto; min-width: 320px; max-width: 640px; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; background-color: transparent;">
<div style="border-collapse: collapse;display: table;width: 100%;background-color:transparent;">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#2e9fdd;"><tr><td align="center"><table cellpadding="0" cellspacing="0" border="0" style="width:640px"><tr class="layout-full-width" style="background-color:transparent"><![endif]-->
<!--[if (mso)|(IE)]><td align="center" width="320" style="background-color:transparent;width:320px; border-top: 0px solid #FFFFFF; border-left: 0px solid #FFFFFF; border-bottom: 0px solid #FFFFFF; border-right: 0px solid #FFFFFF;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:5px; padding-bottom:5px;"><![endif]-->
<div class="col num12" style="min-width: 320px; max-width: 320px; display: table-cell; vertical-align: top; width: 320px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid #FFFFFF; border-left:0px solid #FFFFFF; border-bottom:0px solid #FFFFFF; border-right:0px solid #FFFFFF; padding-top:5px; padding-bottom:5px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 40px; padding-left: 40px; padding-top: 30px; padding-bottom: 0px; font-family: Georgia, serif"><![endif]-->
<div style="color:#ffffff;font-family:Merriwheater, Georgia, serif;line-height:1.2;padding-top:30px;padding-right:40px;padding-bottom:0px;padding-left:40px;">
<div style="line-height: 1.2; font-size: 12px; color: #ffffff; font-family: Merriwheater, Georgia, serif; mso-line-height-alt: 14px;">
<p style="line-height: 1.2; word-break: break-word; text-align: right; mso-line-height-alt: 14px; margin: 0;"><span style="color: #000000;"><span style="font-size: 30px;"><strong>Download our app for free</strong></span></span></p>
</div>
</div>
<!--[if mso]></td></tr></table><![endif]-->
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 40px; padding-left: 40px; padding-top: 8px; padding-bottom: 0px; font-family: Georgia, serif"><![endif]-->
<div style="color:#555555;font-family:Merriwheater, Georgia, serif;line-height:1.8;padding-top:8px;padding-right:40px;padding-bottom:0px;padding-left:40px;">
<div style="font-size: 14px; line-height: 1.8; font-family: Merriwheater, Georgia, serif; color: #555555; mso-line-height-alt: 25px;">
<p style="font-size: 16px; line-height: 1.8; word-break: break-word; text-align: right; font-family: Merriwheater, Georgia, serif; mso-line-height-alt: 29px; margin: 0;"><span style="color: #000000; font-size: 16px;">Grapido: A place where you are more than a user. Grapido is an organic ecosphere of boundless prospects: a platform that generates and amplifies your reach. Connect, collaborate, network and create! Grapido’s powerful algorithms serve to find you the best events and connections.</span></p>
</div>
</div>
<!--[if mso]></td></tr></table><![endif]-->
<div align="right" class="button-container" style="padding-top:10px;padding-right:40px;padding-bottom:5px;padding-left:40px;">
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="border-spacing: 0; border-collapse: collapse; mso-table-lspace:0pt; mso-table-rspace:0pt;"><tr><td style="padding-top: 10px; padding-right: 40px; padding-bottom: 5px; padding-left: 40px" align="right"><v:roundrect xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w="urn:schemas-microsoft-com:office:word" href="" style="height:28.5pt; width:135.75pt; v-text-anchor:middle;" arcsize="90%" strokeweight="0.75pt" strokecolor="#2e9fdd" fillcolor="#000000"><w:anchorlock/><v:textbox inset="0,0,0,0"><center style="color:#ffffff; font-family:Tahoma, Verdana, sans-serif; font-size:18px"><![endif]-->
<a href="https://play.google.com/store/apps/details?id=com.inceptiowave.android.grapido&hl=en_IN"><div style="text-decoration:none;display:inline-block;color:#ffffff;background-color:#000000;border-radius:34px;-webkit-border-radius:34px;-moz-border-radius:34px;width:auto; width:auto;;border-top:1px solid #2e9fdd;border-right:1px solid #2e9fdd;border-bottom:1px solid #2e9fdd;border-left:1px solid #2e9fdd;padding-top:0px;padding-bottom:0px;font-family:'Lato', Tahoma, Verdana, Segoe, sans-serif;text-align:center;mso-border-alt:none;word-break:keep-all;"><span style="padding-left:18px;padding-right:18px;font-size:18px;display:inline-block;"><span style="font-size: 16px; margin: 0; line-height: 2; word-break: break-word; font-family: Lato, Tahoma, Verdana, Segoe, sans-serif; mso-line-height-alt: 32px;"><span data-mce-style="font-size: 18px; line-height: 36px;" style="font-size: 18px; line-height: 36px;"><strong>DOWNLOAD</strong></span></span></span></div></a>
<!--[if mso]></center></v:textbox></v:roundrect></td></tr></table><![endif]-->
</div>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td><td align="center" width="320" style="background-color:transparent;width:320px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:5px; padding-bottom:0px;"><![endif]-->
<!-- <div class="col num6" style="min-width: 320px; max-width: 320px; display: table-cell; vertical-align: top; width: 320px;">
<div style="width:100% !important;">
[if (!mso)&(!IE)]><! -->
<!-- <div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:5px; padding-bottom:0px; padding-right: 0px; padding-left: 0px;"> -->
<!--<![endif]-->
<!-- <div align="center" class="img-container center autowidth" style="padding-right: 0px;padding-left: 0px;"> -->
<!-- [if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr style="line-height:0px"><td style="padding-right: 0px;padding-left: 0px;" align="center"><![endif]<img align="center" alt="Alternate text" border="0" class="center autowidth" src="https://srmpediaesyes.files.wordpress.com/2020/08/screenshot_20200808-191357_01_iphone7plussilver_portrait.png" style="text-decoration: none; -ms-interpolation-mode: bicubic; height: auto; border: 0; width: 100%; max-width: 320px; display: block;" title="Alternate text" width="320"/> -->
<!--[if mso]></td></tr></table><![endif]-->
</div>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div> 
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td></tr></table></td></tr></table><![endif]-->
</div>
</div>
</div>


<div style="background-color:transparent;">
<div class="block-grid" style="Margin: 0 auto; min-width: 320px; max-width: 640px; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; background-color: transparent;">
<div style="border-collapse: collapse;display: table;width: 100%;background-color:transparent;">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:transparent;"><tr><td align="center"><table cellpadding="0" cellspacing="0" border="0" style="width:640px"><tr class="layout-full-width" style="background-color:transparent"><![endif]-->
<!--[if (mso)|(IE)]><td align="center" width="640" style="background-color:transparent;width:640px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:0px; padding-bottom:0px;"><![endif]-->
<div class="col num12" style="min-width: 320px; max-width: 640px; display: table-cell; vertical-align: top; width: 640px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:0px; padding-bottom:0px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<table border="0" cellpadding="0" cellspacing="0" class="divider" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td class="divider_inner" style="word-break: break-word; vertical-align: top; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%; padding-top: 0px; padding-right: 0px; padding-bottom: 0px; padding-left: 0px;" valign="top">
<table align="center" border="0" cellpadding="0" cellspacing="0" class="divider_content" height="30" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; border-top: 0px solid transparent; height: 30px; width: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td height="30" style="word-break: break-word; vertical-align: top; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top"><span></span></td>
</tr>
</tbody>
</table>
</td>
</tr>
</tbody>
</table>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td></tr></table></td></tr></table><![endif]-->
</div>
</div>
</div>
<div style="background-color:transparent;">
<div class="block-grid three-up" style="Margin: 0 auto; min-width: 320px; max-width: 640px; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; background-color: transparent;">
<div style="border-collapse: collapse;display: table;width: 100%;background-color:transparent;">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:transparent;"><tr><td align="center"><table cellpadding="0" cellspacing="0" border="0" style="width:640px"><tr class="layout-full-width" style="background-color:transparent"><![endif]-->
<!--[if (mso)|(IE)]><td align="center" width="213" style="background-color:transparent;width:213px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:0px; padding-bottom:5px;background-color:#f8f8f8;"><![endif]-->

<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td><td align="center" width="213" style="background-color:transparent;width:213px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:0px; padding-bottom:5px;background-color:#f8f8f8;"><![endif]-->

<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td><td align="center" width="213" style="background-color:transparent;width:213px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:0px; padding-bottom:0px;background-color:#f8f8f8;"><![endif]-->
<p>
	Hey """+row[1]+"""!
</p>
<p style="line-height: 1.6;">
	It was a great experience working with you and we wish you all a great luck for your life ahead and may you achieve all your goals and ambitions. We've enclosed your certificate in this email.

Team Inception Wave feels grateful to have you as our interns. We seek your support for the future as well. You can follow us on Instagram, Facebook and LinkedIn to show your support!
</p>
<br><br>
<p>Regards,<br><strong></strong>Inception Wave Pvt. Ltd.</strong></p>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td></tr></table></td></tr></table><![endif]-->
</div>
</div>
</div>
<div style="background-color:transparent;">
<div class="block-grid" style="Margin: 0 auto; min-width: 320px; max-width: 640px; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; background-color: transparent;">
<div style="border-collapse: collapse;display: table;width: 100%;background-color:transparent;">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:transparent;"><tr><td align="center"><table cellpadding="0" cellspacing="0" border="0" style="width:640px"><tr class="layout-full-width" style="background-color:transparent"><![endif]-->
<!--[if (mso)|(IE)]><td align="center" width="640" style="background-color:transparent;width:640px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:0px; padding-bottom:0px;"><![endif]-->
<div class="col num12" style="min-width: 320px; max-width: 640px; display: table-cell; vertical-align: top; width: 640px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:0px; padding-bottom:0px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<table border="0" cellpadding="0" cellspacing="0" class="divider" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td class="divider_inner" style="word-break: break-word; vertical-align: top; min-width: 100%; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%; padding-top: 0px; padding-right: 0px; padding-bottom: 0px; padding-left: 0px;" valign="top">
<table align="center" border="0" cellpadding="0" cellspacing="0" class="divider_content" height="40" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt; border-top: 0px solid transparent; height: 40px; width: 100%;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td height="40" style="word-break: break-word; vertical-align: top; -ms-text-size-adjust: 100%; -webkit-text-size-adjust: 100%;" valign="top"><span></span></td>
</tr>
</tbody>
</table>
</td>
</tr>
</tbody>
</table>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td></tr></table></td></tr></table><![endif]-->
</div>
</div>
</div>
<div style="background-color:#000000;">
<div class="block-grid mixed-two-up" style="Margin: 0 auto; min-width: 320px; max-width: 640px; overflow-wrap: break-word; word-wrap: break-word; word-break: break-word; background-color: transparent;">
<div style="border-collapse: collapse;display: table;width: 100%;background-color:transparent;">
<!--[if (mso)|(IE)]><table width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#000000;"><tr><td align="center"><table cellpadding="0" cellspacing="0" border="0" style="width:640px"><tr class="layout-full-width" style="background-color:transparent"><![endif]-->
<!--[if (mso)|(IE)]><td align="center" width="426" style="background-color:transparent;width:426px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:5px; padding-bottom:5px;"><![endif]-->
<div class="col num8" style="display: table-cell; vertical-align: top; min-width: 320px; max-width: 424px; width: 426px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:5px; padding-bottom:5px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 20px; padding-left: 20px; padding-top: 10px; padding-bottom: 10px; font-family: serif"><![endif]-->
<div style="color:#555555;font-family:'Merriwheater', 'Georgia', serif;line-height:1.2;padding-top:10px;padding-right:20px;padding-bottom:10px;padding-left:20px;">
<div style="line-height: 1.2; font-size: 12px; font-family: 'Merriwheater', 'Georgia', serif; color: #555555; mso-line-height-alt: 14px;">
<p style="font-size: 14px; line-height: 1.2; word-break: break-word; text-align: left; font-family: Merriwheater, Georgia, serif; mso-line-height-alt: 17px; margin: 0;"><em><span style="font-size: 18px; color: #2e9fdd;"><span style="">Contact us:</span></span></em></p>
</div>
</div>
<!--[if mso]></td></tr></table><![endif]-->
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 20px; padding-left: 20px; padding-top: 0px; padding-bottom: 10px; font-family: Tahoma, Verdana, sans-serif"><![endif]-->
<div style="color:#555555;font-family:'Lato', Tahoma, Verdana, Segoe, sans-serif;line-height:1.2;padding-top:0px;padding-right:20px;padding-bottom:10px;padding-left:20px;">
<div style="line-height: 1.2; font-size: 12px; font-family: 'Lato', Tahoma, Verdana, Segoe, sans-serif; color: #555555; mso-line-height-alt: 14px;">
<p style="font-size: 14px; line-height: 1.2; word-break: break-word; text-align: left; font-family: Lato, Tahoma, Verdana, Segoe, sans-serif; mso-line-height-alt: 17px; margin: 0;"><span style="color: #ffffff;"><a style="color: white;" href="mailto:wave.inception@gmail.com">wave.inception@gmail.com</a></span></p>
</div>
</div>
<!--[if mso]></td></tr></table><![endif]-->
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 20px; padding-left: 20px; padding-top: 10px; padding-bottom: 10px; font-family: Tahoma, Verdana, sans-serif"><![endif]-->
<div style="color:#555555;font-family:'Lato', Tahoma, Verdana, Segoe, sans-serif;line-height:1.2;padding-top:10px;padding-right:20px;padding-bottom:10px;padding-left:20px;">
<div style="line-height: 1.2; font-size: 12px; font-family: 'Lato', Tahoma, Verdana, Segoe, sans-serif; color: #555555; mso-line-height-alt: 14px;">
<p style="font-size: 14px; line-height: 1.2; word-break: break-word; text-align: left; font-family: Lato, Tahoma, Verdana, Segoe, sans-serif; mso-line-height-alt: 17px; margin: 0;"><span style="color: #ffffff;">Ameliorating Work Culture </span></p>
</div>
</div>
<!--[if mso]></td></tr></table><![endif]-->
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td><td align="center" width="213" style="background-color:transparent;width:213px; border-top: 0px solid transparent; border-left: 0px solid transparent; border-bottom: 0px solid transparent; border-right: 0px solid transparent;" valign="top"><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr><td style="padding-right: 0px; padding-left: 0px; padding-top:5px; padding-bottom:5px;"><![endif]-->
<div class="col num4" style="display: table-cell; vertical-align: top; max-width: 320px; min-width: 212px; width: 213px;">
<div style="width:100% !important;">
<!--[if (!mso)&(!IE)]><!-->
<div style="border-top:0px solid transparent; border-left:0px solid transparent; border-bottom:0px solid transparent; border-right:0px solid transparent; padding-top:5px; padding-bottom:5px; padding-right: 0px; padding-left: 0px;">
<!--<![endif]-->
<div align="left" class="img-container left autowidth" style="padding-right: 20px;padding-left: 10px;">
<!--[if mso]><table width="100%" cellpadding="0" cellspacing="0" border="0"><tr style="line-height:0px"><td style="padding-right: 20px;padding-left: 10px;" align="left"><![endif]-->
<div style="font-size:1px;line-height:10px"> </div>
<!--[if mso]></td></tr></table><![endif]-->
</div>
<table cellpadding="0" cellspacing="0" class="social_icons" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-lspace: 0pt; mso-table-rspace: 0pt;" valign="top" width="100%">
<tbody>
<tr style="vertical-align: top;" valign="top">
<td style="word-break: break-word; vertical-align: top; padding-top: 10px; padding-right: 10px; padding-bottom: 10px; padding-left: 10px;" valign="top">
<table align="left" cellpadding="0" cellspacing="0" class="social_table" role="presentation" style="table-layout: fixed; vertical-align: top; border-spacing: 0; border-collapse: collapse; mso-table-tspace: 0; mso-table-rspace: 0; mso-table-bspace: 0; mso-table-lspace: 0;" valign="top">
<tbody>
<tr align="center" style="vertical-align: top; display: inline-block; text-align: center;" valign="top">
<td style="word-break: break-word; vertical-align: top; padding-bottom: 0; padding-right: 7.5px; padding-left: 0px;" valign="top"><a href="https://www.instagram.com/inception_wave/" target="_blank"><img alt="Instagram" height="32" src="https://www.freepnglogos.com/uploads/instagram-logos-png-images-free-download-2.png" style="text-decoration: none; -ms-interpolation-mode: bicubic; height: auto; border: 0; display: block;" title="Instagram" width="32"/></a></td>
<td style="word-break: break-word; vertical-align: top; padding-bottom: 0; padding-right: 7.5px; padding-left: 0px;" valign="top"><a href="https://www.linkedin.com/company/14584806/" target="_blank"><img alt="LinkedIn" height="32" src="https://www.freepnglogos.com/uploads/linkedin-social-media-logo-7.png" style="text-decoration: none; -ms-interpolation-mode: bicubic; height: auto; border: 0; display: block;" title="LinkedIn" width="32"/></a></td>
<td style="word-break: break-word; vertical-align: top; padding-bottom: 0; padding-right: 7.5px; padding-left: 0px;" valign="top"><a href="https://www.facebook.com/inceptionwave1202/" target="_blank"><img alt="Facebook" height="32" src="https://www.freepnglogos.com/uploads/facebook-icons/gallery-white-facebook-icon-10.png" style="text-decoration: none; -ms-interpolation-mode: bicubic; height: auto; border: 0; display: block;" title="Facebook" width="32"/></a></td>
</tr>
</tbody>
</table>
</td>
</tr>
</tbody>
</table>
<!--[if (!mso)&(!IE)]><!-->
</div>
<!--<![endif]-->
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
<!--[if (mso)|(IE)]></td></tr></table></td></tr></table><![endif]-->
</div>
</div>
</div>
<!--[if (mso)|(IE)]></td></tr></table><![endif]-->
</td>
</tr>
</tbody>
</table>
<!--[if (IE)]></div><![endif]-->
</body>
</html>
        """

        # part1 = MIMEText(text, "plain")
        part2 = MIMEText(html, "html")
        # message.attach(part1)
        message.attach(part2)
        

        filename=row[6]
        with open("Final_certi/"+filename, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
        
        encoders.encode_base64(part)
        part.add_header("Content-Disposition",f"attachment; filename= {filename}")

        message.attach(part)


        context = ssl.create_default_context()
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
            server.login(sender_email, password)
            server.sendmail(sender_email, receiver_email, message.as_string())
        count+=1
        print("Certificate Mailed To "+row[0]+"..............",end=' ')
        print(count)
		



		
    



def createDoc(intern):
    document = Document('certi.docx')
    if 'Head' not in document.styles:
        style = document.styles.add_style('Head', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style=document.styles['Head']
    font = style.font
    font.size = Pt(18)
    font.bold=True

    style = document.styles['Normal']
    font = style.font
    font.size = Pt(12)
    font.bold=False
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_before = Pt(18)


    head = document.add_paragraph("TO WHOMSOEVER IT MAY CONCERN")
    head.style=document.styles['Head']
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


    content=[]
    name=""
    for word in re.split('\.|_|-|!| ',intern[1]):
        name+=word.capitalize()
        if len(word)==1:
            name=name+"."
        name=name+" "
    name=name[:-1]
    dept=intern[3]
    gender=intern[2]
    if gender=="M":
        s='m'
    elif gender=="F":
        s='f'
    
    he={'m':"he",'f':"she"}
    his={'m':"his",'f':"her"}
    him={'m':"him",'f':"her"}
    salut={'m':"Mr. ",'f':"Ms. "}
    if dept=="GD":
        designation="Graphic Designer"
    elif dept=="HR":
        designation="Human Resource Intern"
    elif dept=="Marketing":
        designation="Marketing and Research Intern"
    elif dept=="CW":
        designation="Content Writer"

    start=intern[4]
    end=intern[5]


    para1="This is to certify that "+salut[s]+name+" has done "+his[s]+" internship as a "+designation+" at Inception Wave Pvt. Ltd, from "+start+" to "+end+"."
    if dept=="Marketing":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Operations team for our product : Grapido, "+he[s]+" made a valuable contribution towards the market research and digital marketing  ventures of Inception Wave Pvt. Ltd."
    elif dept=="GD":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Design and marketing team for our product : Grapido, "+he[s]+" made a valuable contribution towards the publicity and digital marketing  ventures of Inception Wave Pvt. Ltd."
    elif dept=="HR":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Human Resource team, "+he[s]+" made a valuable contribution towards the overall team management of different domains of Inception Wave Pvt. Ltd."
    elif dept=="CW":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Digital Marketing team for our product : Grapido, "+he[s]+" made a valuable contribution towards the content curation for different campaigns of Inception Wave Pvt. Ltd."

    para3="Throughout the internship, "+his[s]+" efforts and dedication towards the task assigned was praiseworthy. Overall, "+he[s]+" demonstrated good communication skills  with a self-motivated attitude to learn new things. Further, "+his[s]+" performance exceeded the expectations and "+he[s]+" was able to complete the assigned tasks successfully on time. "
    para4="We wish "+him[s]+" all the best for "+his[s]+" future endeavours."

    content.append(para1)
    content.append(para2)
    content.append(para3)
    content.append(para4)

    for i in content:
        para=document.add_paragraph(i)
        para.alignment=WD_ALIGN_PARAGRAPH.LEFT
        para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        para.styles=document.styles['Normal']
    
    empty=document.add_paragraph('Regards,')
    
    stamp=document.add_picture('stamp.png',width=Pt(200))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    count=0
    dir="Response/"
    while os.path.isfile("Response/"+name+".docx"):
        count=count+1
        name=name+"copy "
    filename=name
    document.save(dir+filename+".docx")
    intern.append(filename+".pdf")

    with open('mail.csv', mode='a') as mailfile:
        writer = csv.writer(mailfile, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
        writer.writerow(intern)

def convertPDFwindows():
    wdFormatPDF = 17
    infolder="Response"
    out_folder ="Response_PDF/"
    for in_file_name in os.listdir(infolder):
	    print(in_file_name)
	    in_file=infolder+in_file_name
	    word = comtypes.client.CreateObject('Word.Application')
	    doc = word.Documents.Open(in_file)
	    # print("\n"+in_file+" opened")
	
	    outfile_name=in_file_name.replace("docx","pdf")
	    out_file =out_folder+outfile_name
	    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
	    doc.Close()
	    word.Quit()
	    print("successfully converted"+outfile_name)


def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]

    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    return filename.group(1)


def libreoffice_exec():
    # TODO: Provide support for more platforms
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    return 'libreoffice'

def convertPDF():
    directory = os.fsencode("Response/")
    for file in os.listdir(directory):
     filename = os.fsdecode(file)
     if filename.endswith(".docx"): 
         result = convert_to('Response_PDF/', os.path.join("Response", filename))


def rdCSV():
    with open('Intern_Response.csv', newline='') as f:
        reader = csv.reader(f)
        data = list(reader)
        for intern in data:
            createDoc(intern)

def convert(filename):
    watermark="Response_PDF/"+filename
    pdf_file="watermark.pdf"
    merged_file = "Final_certi/"+filename
    input_file = open(pdf_file,'rb')
    input_pdf = PyPDF2.PdfFileReader(input_file)
    watermark_file = open(watermark,'rb')
    watermark_pdf = PyPDF2.PdfFileReader(watermark_file)
    pdf_page = input_pdf.getPage(0)
    watermark_page = watermark_pdf.getPage(0)
    pdf_page.mergePage(watermark_page)
    output = PyPDF2.PdfFileWriter()
    output.addPage(pdf_page)
    merged_file = open(merged_file,'wb')
    output.write(merged_file)
    merged_file.close()
    watermark_file.close()
    input_file.close()


def watermark():
    directory = os.fsencode("Response_PDF/")
    for file in os.listdir(directory):
     filename = os.fsdecode(file)
     if filename.endswith(".pdf"): 
        #  path=os.path.join("Response_PDF", filename)
         convert(filename)
         





@app.route('/app')
def hello():
    return "Let's Go!"

@app.route('/upload')
def upload_file():
   return render_template('upload.html')

@app.route('/uploader', methods = ['GET', 'POST'])
def upload_f():
    try:
        f=''
        if request.method == 'POST':
            f = request.files['file']
            f.save(secure_filename(f.filename))
    
        if f.filename[-3:]=="csv":
            return render_template('sendmail.html')

      
        else:
            return 'Upload a CSV file!'
    except:
        print("No File uploaded")







@app.route('/sendmail')
def certi_create():
    file=open("mail.csv","w")
    file.close()
    rdCSV()
    convertPDF()
    watermark()
    directory = os.fsencode("Response/")
    for file in os.listdir(directory):
        filename = os.fsdecode(file)
        os.remove(os.path.join("Response", filename))
    directory = os.fsencode("Response_PDF/")
    for file in os.listdir(directory):
        filename = os.fsdecode(file)
        os.remove(os.path.join("Response_PDF", filename))
    
    send_mail()
    return "Success"




if __name__ == '__main__':
    app.run()