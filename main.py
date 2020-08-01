from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import PyPDF2
import sys
import csv
import os.path

import sys
import subprocess
import re
import time




def createDoc(intern):
    document = Document('certi.docx')
    if 'Head' not in document.styles:
        style = document.styles.add_style('Head', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style=document.styles['Head']
    font = style.font
    font.size = Pt(18)
    font.bold=True

    style = document.styles['Normal']
    font = style.font
    font.size = Pt(12)
    font.bold=False
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_before = Pt(18)


    head = document.add_paragraph("TO WHOMSOEVER IT MAY CONCERN")
    head.style=document.styles['Head']
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


    content=[]
    name=""
    for word in intern[1].split():
        name+=word.capitalize()
        name=name+" "

    dept=intern[3]
    gender=intern[2]
    if gender=="M":
        s='m'
    elif gender=="F":
        s='f'
    
    he={'m':"he",'f':"she"}
    his={'m':"his",'f':"her"}
    him={'m':"him",'f':"her"}
    if dept=="GD":
        designation="Graphic Designer"
    elif dept=="HR":
        designation="Human Resource Intern"
    elif dept=="Marketing":
        designation="Marketing and Research Intern"
    elif dept=="CW":
        designation="Content Writer"

    start=intern[4]
    end=intern[5]


    para1="This is to certify that "+name+" has done "+his[s]+" internship as a "+designation+" at Inception Wave Pvt. Ltd, from "+start+" to "+end+"."
    if dept=="Marketing":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Operations team for our product : Grapido, "+he[s]+" made a valuable contribution towards the market research and digital marketing  ventures of Inception Wave Pvt. Ltd."
    elif dept=="GD":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Design and marketing team for our product : Grapido, "+he[s]+" made a valuable contribution towards the publicity and digital marketing  ventures of Inception Wave Pvt. Ltd."
    elif dept=="HR":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Human Resource team, "+he[s]+" made a valuable contribution towards the overall team management of different domains of Inception Wave Pvt. Ltd."
    elif dept=="CW":
        para2="During the internship, "+he[s]+" has closely worked as a part of the Digital Marketing team for our product : Grapido, "+he[s]+" made a valuable contribution towards the content curation for different campaigns of Inception Wave Pvt. Ltd."

    para3="Throughout the internship, "+his[s]+" efforts and dedication towards the task assigned was praiseworthy. Throughout the internship, "+he[s]+" demonstrated good communication skills  with a self-motivated attitude to learn new things. Further, "+his[s]+" performance exceeded the expectations and "+he[s]+" was able to complete the assigned tasks successfully on time. "
    para4="We wish "+him[s]+" all the best for "+his[s]+" future endeavours."

    content.append(para1)
    content.append(para2)
    content.append(para3)
    content.append(para4)

    for i in content:
        para=document.add_paragraph(i)
        para.alignment=WD_ALIGN_PARAGRAPH.LEFT
        para.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        para.styles=document.styles['Normal']
    
    empty=document.add_paragraph('Regards,')
    
    stamp=document.add_picture('stamp.png',width=Pt(200))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    count=0
    while os.path.isfile("Response/"+name+".docx"):
        count=count+1
        name=name+"copy "
    filename="Response/"+name+".docx"
    document.save(filename)

def convertPDFwindows():
    wdFormatPDF = 17
    infolder="Response"
    out_folder ="Response_PDF/"
    for in_file_name in os.listdir(infolder):
	    print(in_file_name)
	    in_file=infolder+in_file_name
	    word = comtypes.client.CreateObject('Word.Application')
	    doc = word.Documents.Open(in_file)
	    # print("\n"+in_file+" opened")
	
	    outfile_name=in_file_name.replace("docx","pdf")
	    out_file =out_folder+outfile_name
	    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
	    doc.Close()
	    word.Quit()
	    print("successfully converted"+outfile_name)


def convert_to(folder, source, timeout=None):
    args = [libreoffice_exec(), '--headless', '--convert-to', 'pdf', '--outdir', folder, source]

    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    return filename.group(1)


def libreoffice_exec():
    # TODO: Provide support for more platforms
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    return 'libreoffice'

def convertPDF():
    directory = os.fsencode("Response/")
    for file in os.listdir(directory):
     filename = os.fsdecode(file)
     if filename.endswith(".docx"): 
         result = convert_to('Response_PDF/', os.path.join("Response", filename))


def rdCSV():
    with open('Intern_Response.csv', newline='') as f:
        reader = csv.reader(f)
        data = list(reader)
        for intern in data:
            createDoc(intern)

def convert(filename):
    watermark="Response_PDF/"+filename
    pdf_file="watermark.pdf"
    merged_file = "Final_certi/"+filename
    input_file = open(pdf_file,'rb')
    input_pdf = PyPDF2.PdfFileReader(input_file)
    watermark_file = open(watermark,'rb')
    watermark_pdf = PyPDF2.PdfFileReader(watermark_file)
    pdf_page = input_pdf.getPage(0)
    watermark_page = watermark_pdf.getPage(0)
    pdf_page.mergePage(watermark_page)
    output = PyPDF2.PdfFileWriter()
    output.addPage(pdf_page)
    merged_file = open(merged_file,'wb')
    output.write(merged_file)
    merged_file.close()
    watermark_file.close()
    input_file.close()


def watermark():
    directory = os.fsencode("Response_PDF/")
    for file in os.listdir(directory):
     filename = os.fsdecode(file)
     if filename.endswith(".pdf"): 
        #  path=os.path.join("Response_PDF", filename)
         convert(filename)
         

def main():
    rdCSV()
    convertPDF()
    watermark()
    directory = os.fsencode("Response/")
    for file in os.listdir(directory):
        filename = os.fsdecode(file)
        os.remove(os.path.join("Response", filename))
    directory = os.fsencode("Response_PDF/")
    for file in os.listdir(directory):
        filename = os.fsdecode(file)
        os.remove(os.path.join("Response_PDF", filename))

    


if __name__=="__main__":
    main()




